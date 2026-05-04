from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from pathlib import Path

doc = Document()

# ── Page margins ──────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1)
section.top_margin  = section.bottom_margin = Inches(1)

NAVY   = RGBColor(0x0D, 0x1B, 0x3E)
GOLD   = RGBColor(0xC9, 0xA8, 0x4C)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF5, 0xF2, 0xEB)
GREEN  = RGBColor(0x27, 0x7A, 0x27)
ORANGE = RGBColor(0xD4, 0x7A, 0x1A)
RED    = RGBColor(0xC0, 0x39, 0x2B)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, sides=('top','bottom','left','right'), size=6, color='CCCCCC'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(size))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def heading(text, level=1, color=None, size=None, center=False):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(size or 22)
        run.font.color.rgb = color or NAVY
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
    elif level == 2:
        run.font.size = Pt(size or 14)
        run.font.color.rgb = color or NAVY
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        run.font.size = Pt(size or 12)
        run.font.color.rgb = color or NAVY
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
    run.font.name = 'Calibri'
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def body(text, indent=False, color=None, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    return p

def bullet(text, sub=False, color=None, bold=False):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    p.style = doc.styles['List Bullet']
    run = p.add_run(('    ' if sub else '') + text)
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    run.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.left_indent = Inches(0.5 if sub else 0.25)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    return p

def divider(color='0D1B3E', thickness=12):
    p = doc.add_paragraph()
    p.clear()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), str(thickness))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)

def spacer(pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(pts)
    p.paragraph_format.space_after  = Pt(0)

def status_table(rows):
    """rows = list of (item, status, notes) tuples. status: DONE / IN PROGRESS / TODO"""
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Set column widths
    widths = [Inches(2.8), Inches(1.2), Inches(2.5)]
    hdr_cells = table.rows[0].cells
    headers = ['Item', 'Status', 'Notes']
    header_colors = ['0D1B3E', '0D1B3E', '0D1B3E']
    for i, (cell, hdr) in enumerate(zip(hdr_cells, headers)):
        cell.width = widths[i]
        set_cell_bg(cell, header_colors[i])
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(hdr)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    status_colors = {'DONE': '27AE60', 'IN PROGRESS': 'E67E22', 'TODO': 'BDC3C7', 'BLOCKED': 'C0392B'}
    status_text_colors = {'DONE': WHITE, 'IN PROGRESS': WHITE, 'TODO': RGBColor(0x44,0x44,0x44), 'BLOCKED': WHITE}

    for item, status, notes in rows:
        row = table.add_row()
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]
        row.cells[2].width = widths[2]

        # Item cell
        set_cell_bg(row.cells[0], 'FAFAFA')
        set_cell_border(row.cells[0])
        p0 = row.cells[0].paragraphs[0]
        p0.clear()
        run = p0.add_run(item)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

        # Status cell
        sc = status_colors.get(status, 'DDDDDD')
        set_cell_bg(row.cells[1], sc)
        set_cell_border(row.cells[1])
        p1 = row.cells[1].paragraphs[0]
        p1.clear()
        run = p1.add_run(status)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = status_text_colors.get(status, RGBColor(0,0,0))
        run.font.name = 'Calibri'
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Notes cell
        set_cell_bg(row.cells[2], 'FAFAFA')
        set_cell_border(row.cells[2])
        p2 = row.cells[2].paragraphs[0]
        p2.clear()
        run = p2.add_run(notes)
        run.font.size = Pt(9.5)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0x44,0x44,0x44)

    spacer(8)

def step_table(steps):
    """steps = list of (num, title, description, owner, deps) tuples"""
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [Inches(0.4), Inches(1.8), Inches(3.2), Inches(1.1)]
    hdr_cells = table.rows[0].cells
    for i, (cell, hdr) in enumerate(zip(hdr_cells, ['#', 'Task', 'What to Do', 'Needs'])):
        cell.width = widths[i]
        set_cell_bg(cell, '0D1B3E')
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(hdr)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (num, title, desc, deps) in enumerate(steps):
        row = table.add_row()
        for j, w in enumerate(widths):
            row.cells[j].width = w
        bg = 'F5F2EB' if i % 2 == 0 else 'FFFFFF'

        # Number
        set_cell_bg(row.cells[0], 'C9A84C')
        set_cell_border(row.cells[0])
        p = row.cells[0].paragraphs[0]
        p.clear()
        run = p.add_run(str(num))
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Title
        set_cell_bg(row.cells[1], bg)
        set_cell_border(row.cells[1])
        p = row.cells[1].paragraphs[0]
        p.clear()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = NAVY

        # Description
        set_cell_bg(row.cells[2], bg)
        set_cell_border(row.cells[2])
        p = row.cells[2].paragraphs[0]
        p.clear()
        run = p.add_run(desc)
        run.font.size = Pt(9.5)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0x33,0x33,0x33)

        # Deps
        set_cell_bg(row.cells[3], bg)
        set_cell_border(row.cells[3])
        p = row.cells[3].paragraphs[0]
        p.clear()
        run = p.add_run(deps)
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0x66,0x66,0x66)
        run.italic = True

    spacer(8)


# ══════════════════════════════════════════════════════════════
#  COVER / TITLE
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.clear()
run = p.add_run('EVERMORE LIFE')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after  = Pt(2)

p2 = doc.add_paragraph()
p2.clear()
run2 = p2.add_run('Project Playbook & Roadmap')
run2.bold = True
run2.font.size = Pt(16)
run2.font.color.rgb = GOLD
run2.font.name = 'Calibri'
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(4)

p3 = doc.add_paragraph()
p3.clear()
run3 = p3.add_run('AI Assistant (Sarah)  |  KVN Agent  |  Sales Materials')
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(0x66,0x66,0x66)
run3.italic = True
run3.font.name = 'Calibri'
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(2)

p4 = doc.add_paragraph()
p4.clear()
run4 = p4.add_run('Last updated: April 29, 2026')
run4.font.size = Pt(9)
run4.font.color.rgb = RGBColor(0x99,0x99,0x99)
run4.font.name = 'Calibri'
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_after = Pt(8)

divider('C9A84C', 18)
spacer(8)


# ══════════════════════════════════════════════════════════════
#  SECTION 1: OVERVIEW
# ══════════════════════════════════════════════════════════════
heading('What We Are Building', 1)
divider('0D1B3E', 8)
body('Three interconnected systems that work together to generate, qualify, and close final expense life insurance leads:')
spacer(4)

bullet('Sarah AI Assistant — A virtual office experience on the Evermore Life website. Visitors chat with Sarah, who qualifies them and books appointments with real licensed agents.', bold=False)
bullet('KVN (Obsidian Kevin Prime) — An internal AI accountability agent for Keenan. Lives on the home computer, knows his schedule, and keeps him on track daily.', bold=False)
bullet('Evermore Life Sales Materials — Agent binder and lead tracker already built. Final expense ad funnel and landing page still to be connected.', bold=False)

spacer(8)

# ══════════════════════════════════════════════════════════════
#  SECTION 2: STATUS SNAPSHOT
# ══════════════════════════════════════════════════════════════
heading('Where We Are Right Now', 1)
divider('0D1B3E', 8)

# -- Sarah
heading('Sarah AI Assistant', 2, color=NAVY)
status_table([
    ('Virtual office entry screen',           'DONE',        'CSS office, Sarah avatar, Evermore branding, welcome card'),
    ('Chat conversation flow',                'DONE',        'Collects topic, age, health, coverage goal, booking intent'),
    ('Appointment booking form',              'DONE',        'Name, phone, email, date/time — 7-day auto-populated calendar'),
    ('Confirmation screen',                   'DONE',        'Personalized with visitor name, date, time, topic'),
    ('Lead data object (leadData)',           'DONE',        'All collected fields ready — logs to console for now'),
    ('Claude / ChatGPT API integration',      'TODO',        'Integration Point 1 marked in code — needs API key'),
    ('GoHighLevel webhook',                   'TODO',        'Integration Point 2 marked — needs GHL webhook URL'),
    ('Calendar booking (GHL or Calendly)',    'TODO',        'Integration Point 3 marked — needs calendar link/API'),
    ('Mobile responsiveness',                 'DONE',        'Sidebar hides on mobile, form stacks vertically'),
    ('Deploy to live URL',                    'TODO',        'Needs hosting — Cloudflare Workers, Vercel, or GHL funnel page'),
])

# -- KVN
heading('KVN — Internal Agent', 2, color=NAVY)
status_table([
    ('3D character (Three.js)',               'DONE',        'Hood, visor, armor, cloak panels, glowing cyan eyes'),
    ('Boot screen + TAP TO ENTER',            'DONE',        'Auto-advances after 4.5s, mobile-safe button size'),
    ('CSS fallback avatar',                   'DONE',        'Shows if WebGL fails on device'),
    ('Mobile tab layout (KVN/CHAT/MISSION)',  'DONE',        'Bottom tab bar, iOS safe area support'),
    ('Keyword-based chat responses',          'DONE',        'Basic response system in JS'),
    ('Claude API — real intelligence',        'TODO',        'Needs Anthropic API key'),
    ('Daily briefing mode',                   'TODO',        'Morning schedule summary feature — not yet built'),
    ('Final expense coaching module',         'TODO',        'KVN knows Evermore products and coaches Keenan — planned'),
    ('Deploy to Cloudflare Workers',          'TODO',        'deploy_kvn.sh ready — run on home Mac when back'),
])

# -- Sales Materials
heading('Evermore Life Sales Materials', 2, color=NAVY)
status_table([
    ('Agent binder (.docx)',                  'DONE',        '9 sections, 13 tables, daily scorecard, call script, objection handling'),
    ('Lead tracker spreadsheet (.xlsx)',      'DONE',        '4 sheets, color-coded status, bar chart, 52 formulas'),
    ('Final expense ad funnel (Facebook/Google)', 'TODO',   'In Evermore Life project folder — needs folder access'),
    ('Landing page',                          'TODO',        'Draft exists in Evermore Life Cowork folder — needs review'),
    ('Sarah deployed and embedded in funnel', 'TODO',        'Depends on Sarah deploy + ad funnel build'),
])

spacer(8)


# ══════════════════════════════════════════════════════════════
#  SECTION 3: NEXT STEPS — SARAH
# ══════════════════════════════════════════════════════════════
heading('Next Steps: Sarah AI Assistant', 1)
divider('0D1B3E', 8)
body('Complete these in order. Each step unlocks the next.', bold=True)
spacer(4)

step_table([
    (1,  'Deploy Sarah to a live URL',
         'Option A (easiest): Add Sarah_Evermore_AI.html to a Cloudflare Worker (same as KVN deploy — copy the pattern from deploy_kvn.sh). Option B: Paste into a GoHighLevel funnel page as a custom HTML block. Option C: Host on Vercel (free, drag-and-drop deploy).',
         'Home Mac or laptop'),
    (2,  'Get Claude API key',
         'Go to console.anthropic.com → API Keys → Create key. Copy it somewhere safe. This unlocks real AI responses for both Sarah AND KVN.',
         'Credit card for billing'),
    (3,  'Wire Sarah to Claude API',
         'In Sarah_Evermore_AI.html, find INTEGRATION POINT 1. Replace the switch/case block with a fetch() call to https://api.anthropic.com/v1/messages. Pass conversationStep + leadData as context. System prompt: "You are Sarah, a warm insurance guide for Evermore Life LLC..."',
         'Step 2 done'),
    (4,  'Set up GoHighLevel account',
         'Create GHL account (or use existing). Create a new Pipeline called "Sarah Leads" with stages: New Lead → Contacted → Appointment Set → Closed. Create a Custom Webhook trigger in Automations.',
         'GHL subscription'),
    (5,  'Connect GHL webhook to Sarah',
         'In Sarah_Evermore_AI.html, find INTEGRATION POINT 2. Paste your GHL webhook URL into the fetch() call. Test with a fake booking — verify the contact appears in GHL with all fields (name, phone, email, topic, age, health, coverage, date, time).',
         'Steps 1 + 4 done'),
    (6,  'Set up appointment calendar',
         'In GHL: Calendar → Create Calendar → "Evermore Life Free Consultation" → 15-min slots. OR use Calendly free tier. Get the embed link or API URL.',
         'Step 4 done'),
    (7,  'Wire calendar to Sarah booking screen',
         'In Sarah_Evermore_AI.html, find INTEGRATION POINT 3. Replace the fake form submit with the GHL Calendar embed or a redirect to Calendly pre-filled with the visitor name/email/phone.',
         'Step 6 done'),
    (8,  'Build GHL automation sequence',
         'In GHL Automations: Trigger = Webhook received → Action 1: Create/update contact → Action 2: Add to "Sarah Leads" pipeline → Action 3: Send SMS confirmation to visitor → Action 4: Send internal notification to Keenan.',
         'Steps 4 + 5 done'),
    (9,  'Test full end-to-end flow',
         'Go through Sarah as a fake prospect. Check: (a) conversation feels natural, (b) booking form submits, (c) GHL contact created, (d) calendar event booked, (e) SMS confirmation sent, (f) Keenan gets notified.',
         'Steps 1-8 done'),
    (10, 'Embed Sarah in Evermore Life website',
         'Add an <iframe> or embed Sarah\'s URL into the Evermore Life website/funnel page. Add a CTA button: "Talk to Sarah — Get Your Free Quote" that opens the chat.',
         'Step 1 done'),
])

spacer(8)


# ══════════════════════════════════════════════════════════════
#  SECTION 4: NEXT STEPS — KVN
# ══════════════════════════════════════════════════════════════
heading('Next Steps: KVN — Internal Agent', 1)
divider('0D1B3E', 8)

step_table([
    (1,  'Deploy KVN to Cloudflare Workers',
         'On home Mac: open Terminal, run the deploy_kvn.sh script in the outputs folder. Takes ~10 seconds. URL will be: https://kvn-agent.newmexicomarketingpartners.workers.dev',
         'Home Mac, Terminal'),
    (2,  'Verify KVN on phone',
         'Open the URL on your phone. Confirm: boot screen loads, TAP TO ENTER works, tabs switch (KVN / CHAT / MISSION), CSS fallback shows if WebGL fails.',
         'Step 1 done'),
    (3,  'Get Claude API key (shared with Sarah)',
         'Same key as Sarah — console.anthropic.com. One key works for both.',
         'Credit card'),
    (4,  'Wire KVN chat to Claude API',
         'In KVN.html, find the keyword-based chat handler. Replace with a fetch() to Anthropic API. System prompt: "You are KVN (Obsidian Kevin Prime), Keenan\'s personal AI accountability agent. You know his business is Evermore Life final expense insurance. You are direct, motivating, and know his goals."',
         'Step 3 done'),
    (5,  'Build daily briefing mode',
         'Add a BRIEFING tab or morning trigger. KVN pulls today\'s calendar events (from Google Calendar API or manually entered) and gives Keenan a priorities rundown each morning.',
         'Step 4 done'),
    (6,  'Add final expense coaching module',
         'Feed KVN knowledge of Evermore Life products (final expense, IUL, term). KVN can quiz Keenan on objection handling, product specs, and carrier options.',
         'Step 4 done'),
])

spacer(8)


# ══════════════════════════════════════════════════════════════
#  SECTION 5: NEXT STEPS — AD FUNNEL
# ══════════════════════════════════════════════════════════════
heading('Next Steps: Final Expense Ad Funnel', 1)
divider('0D1B3E', 8)
body('These steps require access to the Evermore Life project folder on the home Mac.', italic=True, color=ORANGE)
spacer(4)

step_table([
    (1,  'Access Evermore Life project folder',
         'Back on home Mac: open Cowork, select the Evermore Life project folder so Claude can read the existing landing page and funnel files.',
         'Home Mac'),
    (2,  'Review existing landing page design',
         'Open the Claude design that was created. Identify what needs to be updated, what copy needs to be written, and how Sarah will be embedded.',
         'Step 1 done'),
    (3,  'Write final expense ad copy',
         'Facebook ad: headline, primary text, CTA. Google ad: 3 headlines, 2 descriptions. Target: ages 50-75, income $25k-$75k, interested in seniors/retirement. Can do this remotely — no folder access needed.',
         'Can do now'),
    (4,  'Build Facebook campaign structure',
         'In Meta Ads Manager: Campaign = Awareness/Lead Gen → Ad Set targeting final expense audience → Ad creative with Sarah\'s office as the visual. CTA: "Get My Free Quote" → lands on Sarah.',
         'Steps 1-3 done'),
    (5,  'Build Google campaign structure',
         'In Google Ads: Search campaign targeting "final expense insurance [state]", "burial insurance", "life insurance for seniors". CTA lands on Sarah.',
         'Steps 1-3 done'),
    (6,  'Connect ad traffic to Sarah',
         'Landing page URL = Sarah\'s deployed URL. UTM parameters on all ad links so GHL tracks source (Facebook vs Google vs organic).',
         'Sarah deployed + Steps 4-5'),
])

spacer(8)


# ══════════════════════════════════════════════════════════════
#  SECTION 6: MASTER PRIORITY ORDER
# ══════════════════════════════════════════════════════════════
heading('Master Priority Order — Do These First', 1)
divider('C9A84C', 12)
spacer(4)

body('Follow this order to get to a working, revenue-generating system as fast as possible:', bold=True)
spacer(6)

priorities = [
    ('1', 'Run deploy_kvn.sh on home Mac',           'KVN goes live. Takes 10 seconds.',                                    '🏠 Home Mac'),
    ('2', 'Run deploy_kvn.sh for Sarah too',          'Sarah goes live at a real URL.',                                      '🏠 Home Mac'),
    ('3', 'Get Claude API key',                       'Unlocks real AI for both Sarah and KVN.',                             '💳 Credit card'),
    ('4', 'Set up GoHighLevel',                       'CRM backbone — everything else plugs into this.',                     '💳 GHL subscription'),
    ('5', 'Wire GHL webhook to Sarah',                'Leads start flowing into your CRM automatically.',                    'Steps 3 + 4'),
    ('6', 'Write Facebook/Google ad copy',            'Can do this now, remotely, no computer needed.',                      '✅ Do now'),
    ('7', 'Access Evermore Life folder on home Mac',  'Review landing page, finalize funnel.',                               '🏠 Home Mac'),
    ('8', 'Launch first ad campaign',                 'Traffic flows to Sarah → leads flow to GHL → Keenan gets notified.', 'Steps 1-7'),
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
for i, (cell, hdr) in enumerate(zip(table.rows[0].cells, ['Priority', 'Task', 'Why It Matters', 'Prerequisites'])):
    set_cell_bg(cell, '0D1B3E')
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(hdr)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

col_widths = [Inches(0.6), Inches(2.2), Inches(2.8), Inches(1.2)]
for cells in [table.rows[0].cells]:
    for i, w in enumerate(col_widths):
        cells[i].width = w

for idx, (num, task, why, prereq) in enumerate(priorities):
    row = table.add_row()
    bg = 'F5F2EB' if idx % 2 == 0 else 'FFFFFF'
    for i, w in enumerate(col_widths):
        row.cells[i].width = w

    set_cell_bg(row.cells[0], 'C9A84C')
    set_cell_border(row.cells[0])
    p = row.cells[0].paragraphs[0]
    p.clear()
    run = p.add_run(num)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(13)
    run.font.name = 'Calibri'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for j, (content, c) in enumerate([(task, NAVY), (why, RGBColor(0x33,0x33,0x33)), (prereq, RGBColor(0x66,0x66,0x66))], 1):
        set_cell_bg(row.cells[j], bg)
        set_cell_border(row.cells[j])
        p = row.cells[j].paragraphs[0]
        p.clear()
        run = p.add_run(content)
        run.font.size = Pt(9.5 if j < 3 else 9)
        run.font.name = 'Calibri'
        run.font.color.rgb = c
        if j == 1:
            run.bold = True
        if j == 3:
            run.italic = True

spacer(10)


# ══════════════════════════════════════════════════════════════
#  SECTION 7: WHAT YOU CAN DO RIGHT NOW (REMOTELY)
# ══════════════════════════════════════════════════════════════
heading('What You Can Do Right Now (No Home Mac Needed)', 1)
divider('0D1B3E', 8)
body('You\'re traveling. Here\'s what can be done from anywhere:', bold=True)
spacer(4)

bullet('Get Claude API key — console.anthropic.com takes 5 minutes')
bullet('Get GoHighLevel account set up — app.gohighlevel.com')
bullet('Write final expense Facebook and Google ad copy (ask Claude to draft it)')
bullet('Review and approve Sarah\'s conversation flow — test the HTML file on your laptop browser')
bullet('Plan your GHL pipeline stages and automation triggers')
bullet('Decide: Cloudflare Workers for hosting, or embed Sarah directly in a GHL funnel page?')

spacer(10)


# ══════════════════════════════════════════════════════════════
#  SECTION 8: FILE LOCATIONS
# ══════════════════════════════════════════════════════════════
heading('File Locations (on Home Mac)', 1)
divider('0D1B3E', 8)
body('All files are in the Cowork outputs folder. Open Cowork on your home Mac to access them.', italic=True)
spacer(4)

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
for i, (cell, hdr) in enumerate(zip(table2.rows[0].cells, ['File', 'What It Is', 'Status'])):
    set_cell_bg(cell, '0D1B3E')
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(hdr)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

file_rows = [
    ('Sarah_Evermore_AI.html',         'Full Sarah AI assistant — virtual office, chat, booking, confirmation',   'DONE — ready to deploy'),
    ('KVN.html',                       'KVN internal agent — 3D character, boot screen, mobile tabs, chat',       'DONE — run deploy_kvn.sh'),
    ('deploy_kvn.sh',                  'Shell script to deploy KVN to Cloudflare Workers in 10 seconds',          'Ready — run in Terminal'),
    ('Evermore_Life_Agent_Binder.docx','9-section agent binder — scripts, objections, carrier info, IUL guide',   'DONE'),
    ('Evermore_Life_Lead_Tracker.xlsx','Lead tracker, follow-up schedule, daily scorecard, weekly chart',         'DONE'),
]

col_w2 = [Inches(1.8), Inches(3.5), Inches(1.5)]
for row_cells in [table2.rows[0].cells]:
    for i, w in enumerate(col_w2):
        row_cells[i].width = w

for idx, (fname, desc, status) in enumerate(file_rows):
    row = table2.add_row()
    bg = 'F5F2EB' if idx % 2 == 0 else 'FFFFFF'
    for i, w in enumerate(col_w2):
        row.cells[i].width = w

    set_cell_bg(row.cells[0], bg)
    set_cell_border(row.cells[0])
    p = row.cells[0].paragraphs[0]
    p.clear()
    run = p.add_run(fname)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    run.font.color.rgb = NAVY

    set_cell_bg(row.cells[1], bg)
    set_cell_border(row.cells[1])
    p = row.cells[1].paragraphs[0]
    p.clear()
    run = p.add_run(desc)
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x33,0x33,0x33)

    set_cell_bg(row.cells[2], bg)
    set_cell_border(row.cells[2])
    p = row.cells[2].paragraphs[0]
    p.clear()
    run = p.add_run(status)
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    run.font.color.rgb = GREEN
    run.bold = True

spacer(12)

# Footer note
divider('C9A84C', 8)
p = doc.add_paragraph()
p.clear()
run = p.add_run('Evermore Life LLC  |  Simple. Honest. Trusted.  |  Protect What Matters. Leave Peace Behind.')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99,0x99,0x99)
run.italic = True
run.font.name = 'Calibri'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)

# ── Save ──────────────────────────────────────────────────────
PROJECT_ROOT = Path(__file__).resolve().parents[2]
out = PROJECT_ROOT / '03_sales_marketing' / 'playbooks' / 'Evermore_Life_Playbook.docx'
out.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(out))
print(f'Saved: {out}')
